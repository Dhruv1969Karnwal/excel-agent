"""Document inspection and text extraction."""

import asyncio
import os
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List, Optional


async def extract_text_from_docx(file_path: str) -> Dict[str, Any]:
    """
    Extract text and structure from Word document.
    
    Args:
        file_path: Path to .docx file
        
    Returns:
        Dictionary with paragraphs, headings, tables, and metadata
    """
    try:
        from docx import Document
    except ImportError:
        raise ImportError(
            "python-docx not installed. Run: pip install python-docx"
        )
    
    def _extract():
        doc = Document(file_path)
        
        paragraphs = []
        headings = []
        tables = []
        full_text = []
        
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            
            # Detect headings by style
            if para.style and para.style.name and para.style.name.startswith('Heading'):
                try:
                    level = int(para.style.name.replace('Heading ', ''))
                except ValueError:
                    level = 1
                headings.append({"text": text, "level": level})
            
            paragraphs.append(text)
            full_text.append(text)
        
        # Extract tables
        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                table_data.append(row_data)
            if table_data:
                tables.append(table_data)
        
        combined_text = "\n\n".join(full_text)
        
        return {
            "paragraphs": paragraphs,
            "headings": headings,
            "tables": tables,
            "full_text": combined_text,
            "word_count": sum(len(p.split()) for p in paragraphs),
            "paragraph_count": len(paragraphs),
            "table_count": len(tables),
        }
    
    return await asyncio.to_thread(_extract)


async def extract_text_from_pdf(file_path: str) -> Dict[str, Any]:
    """
    Extract text from PDF document.
    
    Args:
        file_path: Path to .pdf file
        
    Returns:
        Dictionary with pages, full text, and metadata
    """
    try:
        import PyPDF2
    except ImportError:
        raise ImportError(
            "PyPDF2 not installed. Run: pip install PyPDF2"
        )
    
    def _extract():
        with open(file_path, 'rb') as file:
            reader = PyPDF2.PdfReader(file)
            
            pages = []
            full_text_parts = []
            
            for page_num, page in enumerate(reader.pages, 1):
                text = page.extract_text() or ""
                pages.append({
                    "page_number": page_num,
                    "text": text,
                    "word_count": len(text.split())
                })
                full_text_parts.append(text)
            
            combined_text = "\n\n".join(full_text_parts)
            
            return {
                "pages": pages,
                "full_text": combined_text,
                "page_count": len(reader.pages),
                "word_count": len(combined_text.split()),
            }
    
    return await asyncio.to_thread(_extract)


async def extract_text_from_txt(file_path: str) -> Dict[str, Any]:
    """
    Extract text from plain text or markdown file.
    
    Args:
        file_path: Path to .txt or .md file
        
    Returns:
        Dictionary with content and metadata
    """
    def _extract():
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            content = f.read()
        
        lines = content.split('\n')
        
        # Try to detect headings in markdown
        headings = []
        for line in lines:
            if line.startswith('#'):
                # Count heading level
                level = len(line) - len(line.lstrip('#'))
                text = line.lstrip('#').strip()
                if text:
                    headings.append({"text": text, "level": level})
        
        return {
            "full_text": content,
            "headings": headings,
            "line_count": len(lines),
            "word_count": len(content.split()),
            "char_count": len(content),
        }
    
    return await asyncio.to_thread(_extract)


async def inspect_document(file_path: str) -> Dict[str, Any]:
    """
    Full inspection of a document file.
    
    Supports: .docx, .doc, .pdf, .txt, .md
    
    Uses full-context approach (no chunking/RAG) - 
    the entire document content is passed to the LLM.
    
    Args:
        file_path: Path to the document
        
    Returns:
        Structured data context for the agent
    """
    ext = os.path.splitext(file_path)[1].lower()
    file_name = Path(file_path).name
    
    # Extract based on file type
    if ext in ['.docx', '.doc']:
        extraction = await extract_text_from_docx(file_path)
        doc_type = "Word Document"
    elif ext == '.pdf':
        extraction = await extract_text_from_pdf(file_path)
        doc_type = "PDF Document"
    elif ext in ['.txt', '.md']:
        extraction = await extract_text_from_txt(file_path)
        doc_type = "Text File" if ext == '.txt' else "Markdown File"
    else:
        raise ValueError(f"Unsupported document type: {ext}")
    
    # Generate human-readable description
    description_parts = [
        "Document Overview:",
        f"- File: {file_name}",
        f"- Type: {doc_type}",
        f"- Word Count: {extraction.get('word_count', 'N/A')}",
    ]
    
    if 'page_count' in extraction:
        description_parts.append(f"- Pages: {extraction['page_count']}")
    if 'paragraph_count' in extraction:
        description_parts.append(f"- Paragraphs: {extraction['paragraph_count']}")
    if 'table_count' in extraction:
        description_parts.append(f"- Tables: {extraction['table_count']}")
    
    # Add document structure if headings exist
    headings = extraction.get('headings', [])
    if headings:
        description_parts.append("\nDocument Structure:")
        for h in headings[:15]:  # Limit to first 15 headings
            indent = "  " * (h.get('level', 1) - 1)
            description_parts.append(f"{indent}- {h['text']}")
        if len(headings) > 15:
            description_parts.append(f"  ... and {len(headings) - 15} more sections")
    
    # Add content preview
    full_text = extraction.get('full_text', '')
    if full_text:
        preview_length = min(2000, len(full_text))
        description_parts.append(f"\nContent Preview (first {preview_length} chars):")
        description_parts.append(full_text[:preview_length])
        if len(full_text) > preview_length:
            description_parts.append("...")
    
    description = "\n".join(description_parts)
    
    return {
        "file_path": os.path.abspath(file_path),
        "file_name": file_name,
        "document_type": doc_type,
        "analyzed_at": datetime.now().isoformat(),
        "description": description,
        "full_text": full_text,  # Full content for LLM context
        "summary": {
            "word_count": extraction.get('word_count', 0),
            "page_count": extraction.get('page_count'),
            "paragraph_count": extraction.get('paragraph_count'),
            "table_count": extraction.get('table_count'),
            "headings": headings,
        },
        "tables": extraction.get('tables', []),
        "capabilities": [
            "text_search",
            "summarization",
            "question_answering",
            "information_extraction",
            "content_analysis",
        ]
    }
